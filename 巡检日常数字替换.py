import os
import re
import random
import calendar
from datetime import datetime, timedelta
from docx import Document

# 识别工作日库: pip install chinese_calendar
try:
    from chinese_calendar import is_workday
    HAS_CALENDAR_LIB = True
except ImportError:
    HAS_CALENDAR_LIB = False

# ================= 配置中心 =================
TEMPLATE_FILE = "月报模板.docx"
START_DATE = "2024-8-01"  # 也是序号计算的起始点
END_DATE = "2025-10-23"
OUTPUT_FOLDER = "设备月报库"

# 模式设定：每个月最后一个工作日生成月报
MODE = "MONTHLY_LAST_WORKDAY" 
# ============================================

class UniversalGenerator:
    def __init__(self, template_path):
        self.template_path = template_path

    def _get_random_value(self, match):
        try:
            content = match.group(1)
            start_str, end_str = content.split('-')
            if '.' in start_str:
                precision = len(start_str.split('.')[1])
                val = random.uniform(float(start_str), float(end_str))
                return f"{val:.{precision}f}"
            else:
                return str(random.randint(int(start_str), int(end_str)))
        except: return match.group(0)

    def _process_item(self, item, data_map):
        paras = item.paragraphs if hasattr(item, 'paragraphs') else [item]
        for p in paras:
            if "{{" in p.text:
                full_text = "".join(run.text for run in p.runs)
                full_text = re.sub(r"\{\{random:([\d\.-]+)\}\}", self._get_random_value, full_text)
                for key, value in data_map.items():
                    full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
                
                if p.runs:
                    p.runs[0].text = full_text
                    for r in p.runs[1:]: r.text = ""
                else: p.text = full_text

    def generate(self, data_map, output_path):
        if not os.path.exists(self.template_path):
            print(f"错误: 无法找到模板文件 {self.template_path}")
            return
        doc = Document(self.template_path)
        for para in doc.paragraphs: self._process_item(para, data_map)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells: self._process_item(cell, data_map)
        doc.save(output_path)

def is_actually_workday(dt):
    if HAS_CALENDAR_LIB:
        return is_workday(dt)
    return dt.weekday() < 5

def get_month_workday_count(year, month):
    last_day = calendar.monthrange(year, month)[1]
    return sum(1 for d in range(1, last_day + 1) if is_actually_workday(datetime(year, month, d)))

def get_target_dates(start_dt, end_dt):
    target_dates = []
    temp_month = datetime(start_dt.year, start_dt.month, 1)
    while temp_month <= end_dt:
        year, month = temp_month.year, temp_month.month
        last_day_num = calendar.monthrange(year, month)[1]
        check_dt = datetime(year, month, last_day_num)
        while not is_actually_workday(check_dt):
            check_dt -= timedelta(days=1)
        if start_dt <= check_dt <= end_dt:
            target_dates.append(check_dt)
        nm, ny = (month + 1, year) if month < 12 else (1, year + 1)
        temp_month = datetime(ny, nm, 1)
    return target_dates

def main():
    if not os.path.exists(OUTPUT_FOLDER): os.makedirs(OUTPUT_FOLDER)
    gen = UniversalGenerator(TEMPLATE_FILE)
    
    start_dt = datetime.strptime(START_DATE, "%Y-%m-%d")
    end_dt = datetime.strptime(END_DATE, "%Y-%m-%d")
    
    task_dates = get_target_dates(start_dt, end_dt)
    print(f"已启动：从 {START_DATE} 开始计算序号...")

    for dt in task_dates:
        # --- 核心修改：计算从2024年8月开始的序号 ---
        # 逻辑：(当前年-起始年)*12 + (当前月-起始月) + 1
        month_seq = (dt.year - start_dt.year) * 12 + (dt.month - start_dt.month) + 1
        
        workdays_total = get_month_workday_count(dt.year, dt.month)
        
        context = {
            "date": dt.strftime("%Y年%m月%d日"),
            "year": dt.year,
            "month": dt.strftime("%m"),
            "month_index": month_seq,        # 这里现在是序号：8月是1, 9月是2...
            "day": dt.day,
            "weekday": ["星期一","星期二","星期三","星期四","星期五","星期六","星期日"][dt.weekday()],
            "workday_count": workdays_total
        }
        
        file_name = f"{dt.year}年{context['month']}月网络安全月报.docx"   // 文件命名
        gen.generate(context, os.path.join(OUTPUT_FOLDER, file_name))
        print(f"生成：{file_name} | 序号: {month_seq} | 工作日: {workdays_total}天")

if __name__ == "__main__":

    main()

